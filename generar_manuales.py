"""
Generador de Manuales — Agente de Correos (AquaShield)
Genera 2 documentos Word (.docx):
  - Manual_Usuario_AgenteCorreos.docx
  - Manual_Administrador_AgenteCorreos.docx
"""

import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colores corporativos ─────────────────────────────────────
AZUL = RGBColor(68, 85, 99)
AZUL_OSCURO = RGBColor(45, 58, 69)
NARANJA = RGBColor(235, 95, 10)
ARENA = RGBColor(180, 162, 141)
GRIS = RGBColor(143, 156, 167)
NEGRO_SUAVE = RGBColor(74, 74, 74)
BLANCO = RGBColor(255, 255, 255)

# ── Rutas ─────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(SCRIPT_DIR, "static", "logo_aquachile_dark.png")


# ═══════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(paragraph):
    """Agrega campo PAGE al párrafo."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)


def h1(doc, text):
    """Heading 1 — Azul corporativo, 18pt."""
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.color.rgb = AZUL
    run.font.bold = True
    run.font.name = "Calibri"
    return p


def h2(doc, text):
    """Heading 2 — Naranja, 14pt."""
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = NARANJA
    run.font.bold = True
    run.font.name = "Calibri"
    return p


def h3(doc, text):
    """Heading 3 — Azul, 12pt."""
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = AZUL
    run.font.bold = True
    run.font.name = "Calibri"
    return p


def body(doc, text):
    """Párrafo de texto normal — Negro suave, 10pt."""
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = NEGRO_SUAVE
    run.font.name = "Calibri"
    return p


def bullet(doc, text, bold_prefix=None):
    """Bullet point con prefijo bold opcional."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = NEGRO_SUAVE
        r1.font.name = "Calibri"
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = NEGRO_SUAVE
    r2.font.name = "Calibri"
    return p


def note(doc, text):
    """Nota/Tip — Arena itálica, 9pt."""
    p = doc.add_paragraph()
    run = p.add_run("💡 " + text)
    run.font.size = Pt(9)
    run.font.color.rgb = ARENA
    run.font.italic = True
    run.font.name = "Calibri"
    return p


def warning(doc, text):
    """Advertencia — Naranja bold, 9pt."""
    p = doc.add_paragraph()
    run = p.add_run("⚠ " + text)
    run.font.size = Pt(9)
    run.font.color.rgb = NARANJA
    run.font.bold = True
    run.font.name = "Calibri"
    return p


def add_styled_table(doc, headers, rows):
    """Tabla con header azul y filas alternas arena."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = BLANCO
        run.font.name = "Calibri"
        set_cell_shading(cell, "445563")

    # Rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.color.rgb = NEGRO_SUAVE
            run.font.name = "Calibri"
            if row_idx % 2 == 1:
                set_cell_shading(cell, "F5F3F0")

    doc.add_paragraph()  # Espacio después


# ═══════════════════════════════════════════════════════════════
# ESTRUCTURA DEL DOCUMENTO
# ═══════════════════════════════════════════════════════════════

def create_cover(doc, manual_type):
    """Portada premium con logo, línea naranja y metadata."""
    # Espaciado superior
    for _ in range(3):
        doc.add_paragraph()

    # Logo centrado
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Cm(6))

    doc.add_paragraph()

    # Línea decorativa naranja
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = NARANJA
    run.font.size = Pt(10)

    doc.add_paragraph()

    # AQUASHIELD
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AQUASHIELD")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = AZUL
    run.font.name = "Calibri"

    # Nombre del módulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Módulo Correos")
    run.font.size = Pt(18)
    run.font.color.rgb = ARENA
    run.font.name = "Calibri"

    doc.add_paragraph()

    # Tipo de manual
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(manual_type)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = AZUL_OSCURO
    run.font.name = "Calibri"

    doc.add_paragraph()
    doc.add_paragraph()

    # Metadata
    metadata = [
        ("Versión", "1.0"),
        ("Fecha", "Junio 2026"),
        ("Clasificación", "Confidencial — Uso Interno"),
        ("Equipo", "AquaShield Team"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}: ")
        r1.font.size = Pt(9)
        r1.font.color.rgb = GRIS
        r1.font.name = "Calibri"
        r2 = p.add_run(value)
        r2.font.size = Pt(9)
        r2.font.color.rgb = AZUL
        r2.font.bold = True
        r2.font.name = "Calibri"

    # Salto de página
    doc.add_page_break()


def add_header_footer(doc, title):
    """Header con logo + título y footer con paginación."""
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if os.path.exists(LOGO_PATH):
        run = hp.add_run()
        run.add_picture(LOGO_PATH, height=Cm(0.9))

    run = hp.add_run(f"   {title}")
    run.font.size = Pt(8)
    run.font.color.rgb = AZUL
    run.font.name = "Calibri"

    # Línea separadora en header
    p2 = header.add_paragraph()
    run2 = p2.add_run("━" * 80)
    run2.font.size = Pt(5)
    run2.font.color.rgb = ARENA

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("AquaShield by AquaChile  |  Confidencial  |  Pág. ")
    r1.font.size = Pt(7)
    r1.font.color.rgb = GRIS
    r1.font.name = "Calibri"
    add_page_number(fp)


def add_indice(doc, secciones):
    """Índice con lista numerada."""
    h1(doc, "Índice")
    for i, sec in enumerate(secciones, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {sec}")
        run.font.size = Pt(11)
        run.font.color.rgb = AZUL
        run.font.name = "Calibri"
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# SECCIONES DE CONTENIDO
# ═══════════════════════════════════════════════════════════════

def sec_que_es(doc):
    h1(doc, "1. ¿Qué es el Agente de Correos?")
    body(doc, "El Agente de Correos es un módulo de AquaShield que permite generar correos electrónicos de forma instantánea para clientes previamente registrados en una base de datos Excel.")
    body(doc, "Con un solo clic o arrastrando un documento, el agente abre automáticamente Outlook con el destinatario, asunto, cuerpo y adjuntos ya configurados, ahorrando tiempo y reduciendo errores manuales.")
    doc.add_paragraph()
    h2(doc, "Características principales")
    bullet(doc, "Repositorio de clientes con búsqueda en tiempo real.", "📋")
    bullet(doc, "Generación de correos con plantilla personalizada por cliente.", "✉️")
    bullet(doc, "Arrastrar y soltar archivos para adjuntarlos automáticamente.", "📎")
    bullet(doc, "Soporte para múltiples adjuntos simultáneos.", "📦")
    bullet(doc, "Compatibilidad con Outlook Clásico (COM) y Outlook Nuevo (mailto).", "🔄")
    bullet(doc, "Sistema de favoritos para acceso rápido a clientes frecuentes.", "⭐")
    bullet(doc, "Dashboard con estadísticas de uso (hoy, semana, mes, total).", "📊")
    bullet(doc, "Historial completo de correos generados con exportación a Excel.", "📝")
    bullet(doc, "Modo Claro y Oscuro con identidad visual AquaChile.", "🌙")
    bullet(doc, "Auto-recarga cuando se actualiza el Excel de clientes.", "🔄")
    doc.add_paragraph()


def sec_requisitos(doc):
    h1(doc, "2. Requisitos")
    h2(doc, "Requisitos del sistema")
    bullet(doc, "Windows 10 o superior.")
    bullet(doc, "Python 3.10+ instalado (solo para desarrollo).")
    bullet(doc, "Microsoft Outlook instalado y configurado con una cuenta de correo.")
    bullet(doc, "Conexión a la red local (la aplicación corre en localhost).")
    doc.add_paragraph()
    h2(doc, "Acceso")
    body(doc, "La aplicación se ejecuta localmente en el puerto 5055. Una vez iniciada, se abre automáticamente en el navegador predeterminado.")
    body(doc, "URL de acceso: http://localhost:5055")
    note(doc, "No se requiere login ni contraseña. La aplicación es de uso individual.")
    doc.add_paragraph()


def sec_instalacion(doc):
    h1(doc, "3. Instalación")
    h2(doc, "Paso 1: Clonar el repositorio")
    body(doc, "Abrir una terminal y ejecutar:")
    body(doc, "git clone https://github.com/AquaShield-Team/Agente-Correos.git")
    doc.add_paragraph()
    h2(doc, "Paso 2: Instalar dependencias")
    body(doc, "Navegar a la carpeta del proyecto y ejecutar:")
    body(doc, "pip install -r requirements.txt")
    doc.add_paragraph()
    h2(doc, "Paso 3: Verificar Base_Clientes.xlsx")
    body(doc, "Asegurarse de que el archivo Base_Clientes.xlsx exista en la raíz del proyecto con las columnas:")
    add_styled_table(doc,
        ["Columna", "Descripción", "Ejemplo"],
        [
            ["A (Código)", "Código interno del cliente", "001"],
            ["B (Cliente)", "Nombre del cliente", "Salmones Austral"],
            ["C (Para)", "Email(s) de destino (separados por ;)", "juan@empresa.com"],
            ["D (CC)", "Email(s) en copia (opcional)", "copia@empresa.com"],
            ["E (Asunto)", "Plantilla del asunto", "Documentos [CLIENTE]"],
            ["F (Cuerpo)", "Plantilla del cuerpo (HTML con <br>)", "Estimado, adjunto docs..."],
        ]
    )
    note(doc, "Use [CLIENTE] en el asunto o cuerpo para insertar automáticamente el nombre del cliente.")
    doc.add_paragraph()
    h2(doc, "Paso 4: Ejecutar la aplicación")
    body(doc, "Hacer doble clic en 'Iniciar_Agente_Correos.bat' o ejecutar manualmente:")
    body(doc, "python app.py")
    doc.add_paragraph()


def sec_guia_correos(doc):
    h1(doc, "4. Guía de Uso — Correos")
    h2(doc, "4.1 Buscar un cliente")
    body(doc, "En la vista principal 'Repositorio de Clientes', escriba el nombre o correo del cliente en el campo de búsqueda. La tabla se filtra en tiempo real.")
    doc.add_paragraph()

    h2(doc, "4.2 Generar un correo sin adjuntos")
    body(doc, "Haga clic en el botón de avión de papel (✈) en la columna 'Acción' del cliente deseado. Outlook se abrirá automáticamente con todos los campos pre-llenados.")
    doc.add_paragraph()

    h2(doc, "4.3 Generar un correo con adjuntos (Drag & Drop)")
    body(doc, "Arrastre uno o más archivos desde el Explorador de Windows directamente sobre la fila del cliente. El correo se generará con los archivos adjuntos automáticamente.")
    warning(doc, "En Outlook Nuevo, los adjuntos se guardan en la carpeta 'Adjuntos_AquaShield' del Escritorio. Debe arrastrarlos manualmente al correo.")
    doc.add_paragraph()

    h2(doc, "4.4 Favoritos")
    body(doc, "Haga clic en la estrella (☆) junto al nombre del cliente para marcarlo como favorito. Los favoritos aparecen al inicio de la tabla, separados por una línea divisoria.")
    body(doc, "Los favoritos y no favoritos se ordenan alfabéticamente dentro de su grupo.")
    note(doc, "Los favoritos se guardan en el navegador. Si cambia de navegador o borra los datos, se perderán.")
    doc.add_paragraph()

    h2(doc, "4.5 Refrescar la lista de clientes")
    body(doc, "Haga clic en el botón de flechas circulares (🔄) junto al campo de búsqueda para recargar los datos desde el Excel.")
    note(doc, "La lista también se actualiza automáticamente cada 10 segundos si detecta cambios en el archivo Excel.")
    doc.add_paragraph()


def sec_guia_dashboard(doc):
    h1(doc, "5. Dashboard")
    body(doc, "La vista de Dashboard muestra estadísticas de uso del agente:")
    doc.add_paragraph()
    add_styled_table(doc,
        ["Métrica", "Descripción"],
        [
            ["Hoy", "Correos generados en el día actual"],
            ["Esta Semana", "Correos generados desde el lunes"],
            ["Este Mes", "Correos generados desde el día 1 del mes"],
            ["Total", "Total histórico de correos generados"],
        ]
    )
    body(doc, "Debajo se muestra un ranking de los 3 clientes más frecuentes con barras de progreso.")
    doc.add_paragraph()


def sec_guia_historial(doc):
    h1(doc, "6. Historial")
    body(doc, "La vista de Historial registra todos los correos generados con los siguientes datos:")
    bullet(doc, "Fecha y hora de generación.")
    bullet(doc, "Nombre del cliente.")
    bullet(doc, "Asunto utilizado.")
    bullet(doc, "Archivos adjuntos (si los hubo).")
    bullet(doc, "Modo utilizado (Clásico o Nuevo).")
    doc.add_paragraph()
    h2(doc, "Filtros")
    body(doc, "Puede filtrar el historial por texto (cliente o asunto) y por período (Hoy, Esta semana, Este mes, Todos).")
    doc.add_paragraph()
    h2(doc, "Exportar a Excel")
    body(doc, "Haga clic en el botón 'Exportar Excel' para descargar el historial completo en formato .xlsx.")
    doc.add_paragraph()


def sec_ajustes(doc):
    h1(doc, "7. Ajustes")
    body(doc, "La vista de Ajustes permite configurar el comportamiento de la aplicación:")
    doc.add_paragraph()
    add_styled_table(doc,
        ["Opción", "Descripción"],
        [
            ["Base de Clientes", "Abre el archivo Excel para agregar, editar o eliminar clientes."],
            ["Modo Claro", "Cambia el tema visual de oscuro a claro."],
            ["Outlook Nuevo", "Activa compatibilidad con la versión nueva de Outlook (mailto)."],
        ]
    )
    warning(doc, "Outlook Nuevo no permite adjuntos automáticos ni firma con formato HTML. Los archivos se guardan en la carpeta 'Adjuntos_AquaShield' del Escritorio.")
    doc.add_paragraph()


def sec_modo_oscuro(doc):
    h1(doc, "8. Modo Oscuro / Modo Claro")
    body(doc, "La aplicación soporta dos temas visuales siguiendo la identidad visual de AquaChile:")
    bullet(doc, "El tema se puede cambiar desde Ajustes → Modo Claro.", "Modo Claro:")
    bullet(doc, "El tema se aplica desde Ajustes → desactivar Modo Claro.", "Modo Oscuro:")
    body(doc, "La preferencia se guarda automáticamente en el navegador y se restaura al reiniciar.")
    doc.add_paragraph()


def sec_faq(doc):
    h1(doc, "9. Preguntas Frecuentes")
    h2(doc, "¿Por qué Outlook no adjunta los archivos?")
    body(doc, "Verifique que está usando Outlook Clásico (escritorio). Si usa Outlook Nuevo, active la opción en Ajustes → Outlook Nuevo. Los adjuntos se guardarán en una carpeta del Escritorio para que los arrastre manualmente.")
    doc.add_paragraph()

    h2(doc, "¿Puedo agregar varios destinatarios?")
    body(doc, "Sí. En el Excel, separe los correos con punto y coma (;) en las columnas Para y CC.")
    doc.add_paragraph()

    h2(doc, "¿El correo se envía automáticamente?")
    body(doc, "No. El agente solo abre el borrador en Outlook para que usted lo revise y envíe manualmente.")
    doc.add_paragraph()

    h2(doc, "¿Se pueden adjuntar múltiples archivos?")
    body(doc, "Sí. Seleccione varios archivos del Explorador y arrástrelos todos a la vez sobre la fila del cliente.")
    doc.add_paragraph()

    h2(doc, "¿Dónde se guarda el historial?")
    body(doc, "En el archivo historial.json en la carpeta del proyecto. También puede exportarlo a Excel desde la vista de Historial.")
    doc.add_paragraph()

    h2(doc, "¿Qué pasa si el Excel está abierto?")
    body(doc, "La aplicación genera una copia temporal para leer los datos, por lo que no interfiere con el archivo abierto.")
    doc.add_paragraph()


# ── Secciones SOLO Admin ──────────────────────────────────────

def sec_admin(doc):
    h1(doc, "10. Administración")
    h2(doc, "Gestión de la Base de Clientes")
    body(doc, "La base de datos de clientes se almacena en el archivo Base_Clientes.xlsx ubicado en la raíz del proyecto. Para agregar, editar o eliminar clientes:")
    bullet(doc, "Abra el archivo desde Ajustes → 'Abrir Excel'.")
    bullet(doc, "Edite las filas según la estructura de columnas (Código, Cliente, Para, CC, Asunto, Cuerpo).")
    bullet(doc, "Guarde y cierre el archivo. La aplicación detectará los cambios automáticamente.")
    warning(doc, "No cambie el orden de las columnas ni el nombre de la hoja. El sistema lee por posición.")
    doc.add_paragraph()

    h2(doc, "Variables de plantilla")
    body(doc, "En las columnas Asunto y Cuerpo puede usar variables que se reemplazan automáticamente:")
    add_styled_table(doc,
        ["Variable", "Se reemplaza por"],
        [
            ["[CLIENTE]", "Nombre del cliente (columna B)"],
        ]
    )
    doc.add_paragraph()

    h2(doc, "Historial y auditoría")
    body(doc, "El archivo historial.json registra cada correo generado con timestamp, cliente, destinatario, asunto, adjuntos y modo. Este archivo se puede respaldar o borrar si es necesario reiniciar el conteo.")
    doc.add_paragraph()


def sec_archivos(doc):
    h1(doc, "11. Archivos del Proyecto")
    add_styled_table(doc,
        ["Archivo", "Descripción"],
        [
            ["app.py", "Servidor Flask principal. Contiene todas las rutas API y la lógica COM de Outlook."],
            ["Base_Clientes.xlsx", "Base de datos de clientes con plantillas de correo."],
            ["historial.json", "Registro histórico de correos generados (JSON)."],
            ["requirements.txt", "Dependencias Python del proyecto."],
            ["Iniciar_Agente_Correos.bat", "Script de inicio rápido para Windows."],
            ["LEEME.txt", "Instrucciones de instalación para nuevos usuarios."],
            ["static/app.js", "Lógica del frontend (navegación, drag&drop, favoritos, dashboard)."],
            ["static/styles.css", "Estilos CSS con identidad visual AquaChile (dark/light)."],
            ["templates/index.html", "Template HTML principal con todas las vistas."],
            ["static/favicon.ico", "Ícono de la pestaña del navegador."],
            ["static/manifest.json", "Manifiesto PWA para instalación como aplicación."],
            ["static/sw.js", "Service Worker para funcionalidad offline básica."],
            ["static/logo_aquachile_dark.png", "Logo AquaChile para modo claro."],
            ["static/logo_aquachile_white.png", "Logo AquaChile para modo oscuro."],
        ]
    )
    doc.add_paragraph()


def sec_tecnologias(doc):
    h1(doc, "12. Tecnologías Utilizadas")
    add_styled_table(doc,
        ["Tecnología", "Versión", "Uso"],
        [
            ["Python", "3.10+", "Lenguaje principal del backend"],
            ["Flask", "3.x", "Framework web para servir la API y el frontend"],
            ["openpyxl", "3.x", "Lectura/escritura de archivos Excel"],
            ["pywin32", "306+", "Integración con Outlook vía COM (win32com)"],
            ["HTML5 / CSS3 / JS", "—", "Frontend vanilla (sin frameworks)"],
            ["Font Awesome", "6.4.0", "Iconografía del frontend"],
            ["Inter + Quicksand", "—", "Tipografías corporativas (Google Fonts)"],
        ]
    )
    doc.add_paragraph()


def add_soporte(doc):
    h1(doc, "Soporte")
    body(doc, "Para reportar problemas, solicitar mejoras o recibir capacitación, contacte al equipo AquaShield:")
    doc.add_paragraph()
    add_styled_table(doc,
        ["Campo", "Detalle"],
        [
            ["Equipo", "AquaShield Team"],
            ["Responsable", "Marcelo Ramírez"],
            ["Email", "marcelo.ramirez@aquachile.com"],
            ["Teléfono", "+569 8763 4637"],
        ]
    )
    doc.add_paragraph()

    # Footer decorativo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = ARENA
    run.font.size = Pt(8)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("AquaShield by AquaChile\nCultivos Marinos Chiloé — Puerto Montt, Chile")
    r.font.size = Pt(8)
    r.font.color.rgb = GRIS
    r.font.italic = True
    r.font.name = "Calibri"


# ═══════════════════════════════════════════════════════════════
# GENERADORES PRINCIPALES
# ═══════════════════════════════════════════════════════════════

def generate_user_manual():
    doc = Document()

    # Configurar márgenes
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    create_cover(doc, "Manual de Usuario")
    add_header_footer(doc, "Manual de Usuario | Módulo Correos")

    secciones = [
        "¿Qué es el Agente de Correos?",
        "Requisitos",
        "Instalación",
        "Guía de Uso — Correos",
        "Dashboard",
        "Historial",
        "Ajustes",
        "Modo Oscuro / Modo Claro",
        "Preguntas Frecuentes",
        "Soporte",
    ]
    add_indice(doc, secciones)

    sec_que_es(doc)
    sec_requisitos(doc)
    sec_instalacion(doc)
    sec_guia_correos(doc)
    sec_guia_dashboard(doc)
    sec_guia_historial(doc)
    sec_ajustes(doc)
    sec_modo_oscuro(doc)
    sec_faq(doc)
    add_soporte(doc)

    filename = os.path.join(SCRIPT_DIR, "Manual_Usuario_AgenteCorreos.docx")
    doc.save(filename)
    print(f"✅ Generado: {filename}")
    return filename


def generate_admin_manual():
    doc = Document()

    # Configurar márgenes
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    create_cover(doc, "Manual de Administrador")
    add_header_footer(doc, "Manual de Administrador | Módulo Correos")

    secciones = [
        "¿Qué es el Agente de Correos?",
        "Requisitos",
        "Instalación",
        "Guía de Uso — Correos",
        "Dashboard",
        "Historial",
        "Ajustes",
        "Modo Oscuro / Modo Claro",
        "Preguntas Frecuentes",
        "Administración",
        "Archivos del Proyecto",
        "Tecnologías Utilizadas",
        "Soporte",
    ]
    add_indice(doc, secciones)

    sec_que_es(doc)
    sec_requisitos(doc)
    sec_instalacion(doc)
    sec_guia_correos(doc)
    sec_guia_dashboard(doc)
    sec_guia_historial(doc)
    sec_ajustes(doc)
    sec_modo_oscuro(doc)
    sec_faq(doc)
    sec_admin(doc)
    sec_archivos(doc)
    sec_tecnologias(doc)
    add_soporte(doc)

    filename = os.path.join(SCRIPT_DIR, "Manual_Administrador_AgenteCorreos.docx")
    doc.save(filename)
    print(f"✅ Generado: {filename}")
    return filename


# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("🔧 Generando manuales del Agente de Correos...")
    print()
    generate_user_manual()
    generate_admin_manual()
    print()
    print("🎉 Ambos manuales generados exitosamente.")
